"""
document_handler.py
--------------------
Structure-preserving translation for DOCX, TXT, and (best-effort) PDF files.

Every translate_* function takes a `translator_fn` callable with the signature
    translator_fn(texts: list[str], src_lang: str, tgt_lang: str) -> list[str]
so the real IndicTrans2 engine (translator_engine.TranslationEngine.translate_batch)
and a cheap mock (used in tests) are interchangeable.

DOCX strategy
-------------
python-docx exposes text as a tree of paragraphs -> runs. Structure that matters
(headings, bullet/numbered lists, table layout, alignment, page breaks, images)
lives on the *paragraph* and *table* objects, not on individual runs. So:

  1. Walk every paragraph in the document body, every table cell (recursively,
     since cells can contain nested tables), and every header/footer.
  2. Batch-translate all paragraph texts in one go (fast, and keeps model
     context per call small).
  3. Write the translated text into the paragraph's *first* run (so its font,
     bold/italic/color survive) and blank out any remaining runs.
     Paragraph-level formatting (style, alignment, spacing, numbering) is
     untouched because we never touch the paragraph or style objects.

Known limitation: if a single paragraph mixes multiple runs with different
formatting (e.g. one bold word mid-sentence), that fine-grained formatting
collapses onto the first run's style after translation. Document-level
structure (headings/lists/tables/images/page layout) is preserved exactly.

PDF strategy
------------
PDF has no flow-document structure to preserve -- it's absolute-positioned
glyphs. We extract text blocks with PyMuPDF, translate each block, redact the
original glyphs, and draw the translation back into the same bounding box with
auto-shrunk font size. This is best-effort: translated text is often longer or
shorter than the original, so line breaks can shift. Scanned (image-only) PDFs
are not handled -- that needs OCR, listed as a future improvement.
"""

from __future__ import annotations

import sys
from pathlib import Path
from typing import Callable, List, Optional, Tuple

TranslatorFn = Callable[[List[str], str, str], List[str]]


def _assets_dir() -> Path:
    """
    Resolves the project's bundled `fonts/` directory both when running from
    source and when packaged by PyInstaller (which unpacks data files into a
    temp dir referenced by sys._MEIPASS instead of alongside this .py file).
    """
    base = Path(getattr(sys, "_MEIPASS", Path(__file__).resolve().parent))
    return base / "fonts"


# PyMuPDF's built-in fonts (Helvetica, Times, etc.) only cover Latin script --
# they render Devanagari as boxes/question marks. Hindi and Marathi both use
# Devanagari, so PDF output in either language needs a real Devanagari font
# embedded. Shobhika (IIT Bombay, SIL Open Font License) is bundled in
# fonts/ for exactly this. English output keeps using PyMuPDF's built-in
# Helvetica -- no embedding needed for Latin script.
_DEVANAGARI_FONT_FILE = _assets_dir() / "Shobhika-Regular.otf"


def _font_for_language(lang: str) -> Tuple[str, Optional[str]]:
    """Returns (fontname, fontfile) for use with page.insert_textbox()."""
    if lang in ("Hindi", "Marathi") and _DEVANAGARI_FONT_FILE.exists():
        return "shobhika-devanagari", str(_DEVANAGARI_FONT_FILE)
    return "helv", None


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def _iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs
            for nested_table in cell.tables:
                yield from _iter_table_paragraphs(nested_table)


def _iter_all_paragraphs(doc):
    yield from doc.paragraphs
    for table in doc.tables:
        yield from _iter_table_paragraphs(table)

    for section in doc.sections:
        containers = [
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ]
        for container in containers:
            if container is None:
                continue
            yield from container.paragraphs
            for table in container.tables:
                yield from _iter_table_paragraphs(table)


def _set_paragraph_text(paragraph, new_text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def translate_docx(
    input_path: str,
    output_path: str,
    translator_fn: TranslatorFn,
    src_lang: str,
    tgt_lang: str,
    progress_cb=None,
) -> None:
    from docx import Document  # local import: keeps module importable without python-docx installed

    doc = Document(input_path)
    paragraphs = [p for p in _iter_all_paragraphs(doc) if p.text.strip()]

    if progress_cb:
        progress_cb(f"{Path(input_path).name}: found {len(paragraphs)} text segments")

    texts = [p.text for p in paragraphs]
    translated = translator_fn(texts, src_lang, tgt_lang)

    for paragraph, new_text in zip(paragraphs, translated):
        _set_paragraph_text(paragraph, new_text)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


# ---------------------------------------------------------------------------
# TXT
# ---------------------------------------------------------------------------

def translate_txt(
    input_path: str,
    output_path: str,
    translator_fn: TranslatorFn,
    src_lang: str,
    tgt_lang: str,
    progress_cb=None,
) -> None:
    text = Path(input_path).read_text(encoding="utf-8")
    lines = text.split("\n")

    if progress_cb:
        progress_cb(f"{Path(input_path).name}: found {len(lines)} lines")

    translated_lines = translator_fn(lines, src_lang, tgt_lang)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    Path(output_path).write_text("\n".join(translated_lines), encoding="utf-8")


# ---------------------------------------------------------------------------
# PDF (best-effort)
# ---------------------------------------------------------------------------

def _insert_shrink_to_fit(
    page, rect, text: str, fontname: str = "helv", fontfile: Optional[str] = None,
    start_size: int = 10, min_size: int = 5,
) -> None:
    """
    PyMuPDF's insert_textbox does not auto-fit text to a box: the bounding
    box PyMuPDF reports for the *original* text is usually a point or two too
    tight to hold the *translated* text at the same font size (translations
    are rarely the same length), so a literal reinsertion at the original
    rect silently fails (no text drawn, no error). Pad the box slightly and
    step the font size down until it fits; if even the minimum size overflows,
    insert at the minimum size anyway rather than lose the text entirely.

    fontname/fontfile select which font to embed -- PyMuPDF's built-in fonts
    (the default "helv") have no Devanagari glyphs, so Hindi/Marathi output
    must pass a real font file here (see _font_for_language) or it renders as
    "?" placeholders.
    """
    import fitz

    padded = fitz.Rect(rect.x0 - 1, rect.y0 - 1, rect.x1 + 2, rect.y1 + 4)
    kwargs = {"fontname": fontname, "align": 0}
    if fontfile:
        kwargs["fontfile"] = fontfile
    for size in range(start_size, min_size - 1, -1):
        remaining_space = page.insert_textbox(padded, text, fontsize=size, **kwargs)
        if remaining_space >= 0:
            return
    page.insert_textbox(padded, text, fontsize=min_size, **kwargs)


def _rect_in_any(rect, table_bboxes) -> bool:
    """True if rect's center falls inside any of the given table bounding boxes."""
    cx, cy = (rect.x0 + rect.x1) / 2, (rect.y0 + rect.y1) / 2
    for bbox in table_bboxes:
        if bbox.x0 <= cx <= bbox.x1 and bbox.y0 <= cy <= bbox.y1:
            return True
    return False


def translate_pdf(
    input_path: str,
    output_path: str,
    translator_fn: TranslatorFn,
    src_lang: str,
    tgt_lang: str,
    progress_cb=None,
) -> None:
    try:
        import fitz  # PyMuPDF
    except ImportError as exc:
        raise ImportError(
            "PDF support requires PyMuPDF. Install it with:\n    pip install PyMuPDF"
        ) from exc

    doc = fitz.open(input_path)

    # (page_index, rect, text, is_table_cell) -- is_table_cell controls how
    # each item gets redacted below: table cells use an inset rect so the
    # table's border lines survive redaction, loose paragraph text doesn't
    # need that (there are no lines to protect).
    all_items = []

    for page_index, page in enumerate(doc):
        table_bboxes = []
        try:
            tables = page.find_tables()
        except Exception:
            tables = None  # older PyMuPDF without table support -- degrade gracefully

        if tables is not None:
            for table in tables.tables:
                table_bboxes.append(fitz.Rect(table.bbox))
                grid = table.extract()  # row-major list of lists of cell text
                for row, row_cells in zip(grid, table.rows):
                    for cell_text, cell_bbox in zip(row, row_cells.cells):
                        cell_text = (cell_text or "").strip()
                        if cell_text and cell_bbox is not None:
                            all_items.append((page_index, fitz.Rect(cell_bbox), cell_text, True))

        # Loose (non-table) text blocks -- skip anything that falls inside a
        # detected table's bounding box, since that text was already picked
        # up cell-by-cell above and would otherwise be translated twice.
        for b in page.get_text("blocks"):
            x0, y0, x1, y1, text = b[0], b[1], b[2], b[3], b[4]
            text = text.strip()
            if not text:
                continue
            rect = fitz.Rect(x0, y0, x1, y1)
            if _rect_in_any(rect, table_bboxes):
                continue
            all_items.append((page_index, rect, text, False))

    num_tables = sum(1 for _, _, _, is_cell in all_items if is_cell)
    if progress_cb:
        progress_cb(
            f"{Path(input_path).name}: found {len(all_items)} text segments "
            f"({num_tables} table cells, {len(all_items) - num_tables} other)"
        )

    texts = [item[2] for item in all_items]
    translated = translator_fn(texts, src_lang, tgt_lang)

    # Redact original text first (must apply per page before inserting new
    # text). Table cells are redacted with an inset rect so the redaction
    # only clears the text inside the cell, leaving the cell's border lines
    # (drawn exactly on the cell's bounding box) intact.
    for page_index, rect, _text, is_table_cell in all_items:
        page = doc[page_index]
        if is_table_cell:
            inset = fitz.Rect(rect.x0 + 2, rect.y0 + 2, rect.x1 - 2, rect.y1 - 2)
            if inset.is_valid and not inset.is_empty:
                page.add_redact_annot(inset)
            else:
                page.add_redact_annot(rect)
        else:
            page.add_redact_annot(rect)
    for page in doc:
        page.apply_redactions()

    fontname, fontfile = _font_for_language(tgt_lang)
    if progress_cb and fontfile:
        progress_cb(f"Using embedded Devanagari font for {tgt_lang} output.")

    for (page_index, rect, _original, _is_table_cell), new_text in zip(all_items, translated):
        page = doc[page_index]
        _insert_shrink_to_fit(page, rect, new_text, fontname=fontname, fontfile=fontfile)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    doc.close()


# ---------------------------------------------------------------------------
# Dispatch
# ---------------------------------------------------------------------------

SUPPORTED_EXTENSIONS = {".docx": translate_docx, ".txt": translate_txt, ".pdf": translate_pdf}


def translate_file(
    input_path: str,
    output_path: str,
    translator_fn: TranslatorFn,
    src_lang: str,
    tgt_lang: str,
    progress_cb=None,
) -> None:
    ext = Path(input_path).suffix.lower()
    handler = SUPPORTED_EXTENSIONS.get(ext)
    if handler is None:
        raise ValueError(f"Unsupported file type: {ext}")
    handler(input_path, output_path, translator_fn, src_lang, tgt_lang, progress_cb=progress_cb)


def extract_sample_texts(input_path: str, max_samples: int = 20) -> List[str]:
    """
    Pulls a handful of representative text snippets from a document, used for
    language auto-detection before the real translation pass runs.
    """
    ext = Path(input_path).suffix.lower()

    if ext == ".docx":
        from docx import Document

        doc = Document(input_path)
        texts = [p.text for p in _iter_all_paragraphs(doc) if p.text.strip()]
        return texts[:max_samples]

    if ext == ".txt":
        text = Path(input_path).read_text(encoding="utf-8")
        lines = [line for line in text.split("\n") if line.strip()]
        return lines[:max_samples]

    if ext == ".pdf":
        import fitz

        doc = fitz.open(input_path)
        samples: List[str] = []
        for page in doc:
            for b in page.get_text("blocks"):
                text = b[4].strip()
                if text:
                    samples.append(text)
                if len(samples) >= max_samples:
                    break
            if len(samples) >= max_samples:
                break
        doc.close()
        return samples

    raise ValueError(f"Unsupported file type: {ext}")
